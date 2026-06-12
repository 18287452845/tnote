"""将综合实验 Markdown 讲义转换为格式化的 Word 实验手册"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def setup_styles(doc):
    """配置文档样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.35

    # 标题样式
    heading_configs = [
        ('Heading 1', 22, True, RGBColor(0x1a, 0x56, 0x8e)),
        ('Heading 2', 17, True, RGBColor(0x1a, 0x56, 0x8e)),
        ('Heading 3', 14, True, RGBColor(0x2c, 0x6f, 0xad)),
        ('Heading 4', 12, True, RGBColor(0x3a, 0x87, 0xc2)),
    ]
    for name, size, bold, color in heading_configs:
        s = doc.styles[name]
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = color
        s.font.name = '微软雅黑'
        s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        s.paragraph_format.space_before = Pt(16)
        s.paragraph_format.space_after = Pt(8)

    # 代码块样式
    if 'CodeBlock' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9.5)
        code_style.font.color.rgb = RGBColor(0x2b, 0x2b, 0x2b)
        code_style.paragraph_format.space_before = Pt(2)
        code_style.paragraph_format.space_after = Pt(2)
        code_style.paragraph_format.line_spacing = 1.15
        code_style.paragraph_format.left_indent = Cm(0.5)

    # 提示框样式
    if 'AsideBlock' not in [s.name for s in doc.styles]:
        aside_style = doc.styles.add_style('AsideBlock', WD_STYLE_TYPE.PARAGRAPH)
        aside_style.font.name = '微软雅黑'
        aside_style.font.size = Pt(10)
        aside_style.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)
        aside_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        aside_style.paragraph_format.left_indent = Cm(0.8)
        aside_style.paragraph_format.space_before = Pt(2)
        aside_style.paragraph_format.space_after = Pt(2)
        aside_style.paragraph_format.line_spacing = 1.3


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1a568e')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val.strip())
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f0f4f8')

    doc.add_paragraph('')  # 表后空行


def add_code_block(doc, code_lines):
    """添加代码块（带灰色背景）"""
    # 用表格实现背景色效果
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'f5f5f5')
    cell.text = ''

    for line in code_lines:
        p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2b, 0x2b, 0x2b)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1

    # 设置单元格边距
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '<w:top w:w="80" w:type="dxa"/>'
        '<w:left w:w="120" w:type="dxa"/>'
        '<w:bottom w:w="80" w:type="dxa"/>'
        '<w:right w:w="120" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)

    # 设置表格宽度为 100%
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    doc.add_paragraph('')


def add_aside_block(doc, lines):
    """添加提示框（带左边框样式）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'f8f9fa')
    cell.text = ''

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.3

    # 设置左边框颜色
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:left w:val="single" w:sz="18" w:space="0" w:color="3a87c2"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '<w:top w:w="60" w:type="dxa"/>'
        '<w:left w:w="160" w:type="dxa"/>'
        '<w:bottom w:w="60" w:type="dxa"/>'
        '<w:right w:w="100" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)

    # 设置表格宽度为 100%
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    doc.add_paragraph('')


def add_rich_paragraph(doc, text, style='Normal'):
    """添加包含内联格式的段落（加粗、行内代码）"""
    p = doc.add_paragraph(style=style)
    # 处理加粗和行内代码的混合文本
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
        else:
            run = p.add_run(part)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def parse_and_build(doc, md_text):
    """解析 Markdown 并构建 Word 文档"""
    lines = md_text.split('\n')
    i = 0
    total = len(lines)

    while i < total:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平分割线
        if stripped == '---':
            # 添加一条细线
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="cccccc"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # 提示框 <aside>
        if stripped.startswith('<aside>'):
            aside_lines = []
            i += 1
            while i < total and not lines[i].strip().startswith('</aside>'):
                aline = lines[i].strip()
                if aline:
                    aside_lines.append(aline)
                i += 1
            if i < total:
                i += 1  # skip </aside>
            if aside_lines:
                add_aside_block(doc, aside_lines)
            continue

        # 代码块
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < total and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < total:
                i += 1  # skip closing ```
            if code_lines:
                add_code_block(doc, code_lines)
            continue

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # 清理标题中的 emoji
            title = re.sub(r'[🎯🧭🖥️📸📍💬⚠️✅💡🔧🛡️📝📄⚡🏆📋]', '', title).strip()
            doc.add_heading(title, level=level)
            i += 1
            continue

        # 表格
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < total and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # 跳过分隔行
                rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    # 确保列数匹配
                    while len(cells) < len(headers):
                        cells.append('')
                    rows.append(cells[:len(headers)])
                add_table(doc, headers, rows)
            continue

        # 列表项
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', stripped)
        if list_match:
            indent = len(list_match.group(1)) // 2
            marker = list_match.group(2)
            content = list_match.group(3)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8 + indent * 0.6)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            if marker in ['-', '*']:
                bullet = p.add_run('  •  ')
            else:
                num = marker.rstrip('.')
                bullet = p.add_run(f'  {num}.  ')
            bullet.font.name = '微软雅黑'
            bullet.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 内容中的富文本
            parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
                else:
                    run = p.add_run(part)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 普通段落（支持富文本）
        clean = stripped
        # 跳过纯 emoji 行
        if re.match(r'^[^\w\s]+$', clean):
            i += 1
            continue

        add_rich_paragraph(doc, clean)
        i += 1


def add_cover_page(doc, title, subtitle):
    """添加封面页"""
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for _ in range(3):
        doc.add_paragraph('')

    info_lines = [
        '课程：数据库系统管理与运维',
        '环境：Ubuntu 24.04 LTS + MySQL 8.0 + Navicat',
        '密码：统一使用 123456（课堂环境）',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()


def add_student_info_page(doc):
    """添加学生信息页"""
    doc.add_heading('学生信息', level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    labels = ['姓名', '班级', '学号', '实验日期', '指导教师']
    for i, label in enumerate(labels):
        cell_l = table.rows[i].cells[0]
        cell_l.text = ''
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell_l, 'e8edf2')
        cell_l.width = Cm(4)

        cell_r = table.rows[i].cells[1]
        cell_r.text = ''
        cell_r.width = Cm(12)

    doc.add_paragraph('')
    doc.add_page_break()


def convert_md_to_word(md_path, docx_path, title, subtitle):
    """转换单个 Markdown 文件为 Word"""
    md_text = Path(md_path).read_text(encoding='utf-8')

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    setup_styles(doc)

    # 封面
    add_cover_page(doc, title, subtitle)

    # 学生信息页
    add_student_info_page(doc)

    # 跳过原始 Markdown 的第一个 # 标题行（已作为封面）
    first_heading = re.search(r'^#\s+.*$', md_text, re.MULTILINE)
    if first_heading:
        md_content = md_text[first_heading.end():]
    else:
        md_content = md_text

    # 解析并构建文档
    parse_and_build(doc, md_content)

    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('数据库系统管理与运维 - 实验手册')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.save(str(docx_path))
    print(f'  -> {docx_path.name}')


def main():
    base_dir = Path(r'C:\Users\admin\Documents\tnote\tnote\课程讲义\数据库系统管理和运维')
    output_dir = base_dir
    output_dir.mkdir(exist_ok=True)

    experiments = [
        {
            'md': '综合实验一 MySQL安全管理综合实战.md',
            'docx': '综合实验一 MySQL安全管理综合实战-实验手册.docx',
            'title': '综合实验一',
            'subtitle': 'MySQL 安全管理综合实战',
        },
        {
            'md': '综合实验二 MySQL备份恢复与灾难恢复实战.md',
            'docx': '综合实验二 MySQL备份恢复与灾难恢复实战-实验手册.docx',
            'title': '综合实验二',
            'subtitle': 'MySQL 备份恢复与灾难恢复实战',
        },
        {
            'md': '综合实验三 MySQL主从复制与高可用部署实战.md',
            'docx': '综合实验三 MySQL主从复制与高可用部署实战-实验手册.docx',
            'title': '综合实验三',
            'subtitle': 'MySQL 主从复制与高可用部署实战',
        },
        {
            'md': '综合实验四 MySQL性能优化与数据管理实战.md',
            'docx': '综合实验四 MySQL性能优化与数据管理实战-实验手册.docx',
            'title': '综合实验四',
            'subtitle': 'MySQL 性能优化与数据管理实战',
        },
    ]

    print('开始转换...')
    for exp in experiments:
        md_path = base_dir / exp['md']
        docx_path = output_dir / exp['docx']
        if md_path.exists():
            convert_md_to_word(md_path, docx_path, exp['title'], exp['subtitle'])
        else:
            print(f'  [跳过] {md_path.name} 不存在')
    print('全部完成！')


if __name__ == '__main__':
    main()
